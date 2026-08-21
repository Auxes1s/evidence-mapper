from __future__ import annotations

import json, time
from pathlib import Path
import pytest
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen.canvas import Canvas

from repo_research.ollama import MockBackend
from repo_research.research import _decompose_claim, research

@pytest.fixture()
def corpus(tmp_path: Path) -> Path:
    (tmp_path/"primary.md").write_text("# Signed Resolution\nOn 15 March 2023, the Board approved establishment of the Atlas Unit. Operations began on 1 July 2023.\n",encoding="utf-8")
    (tmp_path/"later-summary.txt").write_text("A retrospective summary incorrectly states that the Atlas Unit began operations in June 2023.\n",encoding="utf-8")
    (tmp_path/"proposal.md").write_text("A proposal dated 2 January 2023 recommends creating an Atlas Unit. This document does not implement the proposal.\n",encoding="utf-8")
    (tmp_path/"unsupported.md").write_text("There is no documentary evidence that Atlas doubled evaluation capacity or accelerated every evaluation.\n",encoding="utf-8")
    (tmp_path/"irrelevant.txt").write_text("Atlas is a map collection. An operations manual for a printer was completed in June 2023.\n",encoding="utf-8")
    doc=Document(); doc.add_heading("Handover",1); doc.add_paragraph("The Atlas Unit handover occurred on 10 December 2024."); doc.save(tmp_path/"handover.docx")
    wb=Workbook(); ws=wb.active; ws.title="Milestones"; ws.append(["Event","Date","Status"]); ws.append(["Atlas funding released","2023-05-20","Approved"]); wb.save(tmp_path/"funding.xlsx")
    pdf=Canvas(str(tmp_path/"minutes.pdf")); pdf.drawString(72,720,"Meeting minutes: Atlas implementation commenced on 20 May 2023."); pdf.showPage(); pdf.drawString(72,720,"The Unit was confirmed operational on 1 July 2023."); pdf.save()
    (tmp_path/"status.py").write_text('OPERATIONAL_DATE = "2023-07-01"\n\ndef unit_status():\n    return "operational"\n',encoding="utf-8")
    return tmp_path

def test_inventory_and_format_provenance(corpus):
    out=research(corpus,mode="inventory",depth="quick",backend=MockBackend())
    fam=out["inventory"]["file_families"]
    assert fam["pdf"] == fam["docx"] == fam["xlsx"] == 1
    ev=research(corpus,"When did Atlas become operational?",mode="evidence",depth="quick",fresh=True,backend=MockBackend())
    assert any(x["source_path"]=="minutes.pdf" and x["page"]==2 for x in ev["evidence"])
    sheet=[x for x in research(corpus,"When was Atlas funding released?",mode="evidence",depth="quick",fresh=True,backend=MockBackend())["evidence"] if x["source_path"]=="funding.xlsx"]
    assert sheet and sheet[0]["sheet"]=="Milestones" and sheet[0]["row_start"]==2
    assert all(x["quote_verified"] for x in ev["evidence"])

def test_contradiction_and_stage_distinction(corpus):
    out=research(corpus,"Atlas operations began on 1 July 2023",mode="contradictions",depth="deep",fresh=True,backend=MockBackend())
    assert any(x["source_path"]=="later-summary.txt" for x in out["evidence"])
    chron=research(corpus,"Reconstruct Atlas establishment and operation",mode="chronology",depth="standard",fresh=True,backend=MockBackend())
    paths={x["source_path"] for x in chron["evidence"]}
    assert "proposal.md" in paths and "primary.md" in paths

def test_persistence_staleness_and_fresh(corpus):
    q="When did Atlas become operational?"
    first=research(corpus,q,mode="evidence",depth="quick",fresh=True,backend=MockBackend())
    assert first["evidence"] and (corpus/".codex-research/evidence.jsonl").exists()
    reused=research(corpus,q,mode="evidence",depth="quick",backend=MockBackend())
    assert reused["reused_evidence"]
    time.sleep(.01); (corpus/"primary.md").write_text("# Revised\nOperations began on 2 July 2023.\n",encoding="utf-8")
    after=research(corpus,q,mode="evidence",depth="quick",backend=MockBackend())
    assert not any(x["source_path"]=="primary.md" for x in after["reused_evidence"])
    challenged=research(corpus,q,mode="evidence",depth="quick",challenge_existing=True,backend=MockBackend())
    assert challenged["reused_evidence"] == []

def test_no_invented_sources_and_compound_partial_support(corpus):
    out=research(corpus,"Atlas created the unit, doubled capacity, accelerated evaluations, and institutionalized evaluation",mode="exhaustive",depth="deep",fresh=True,backend=MockBackend())
    actual={str(p.relative_to(corpus)) for p in corpus.rglob("*") if p.is_file()}
    assert all(x["source_path"] in actual for x in out["evidence"])
    assert len(out["propositions"]) >= 3
    assert not any("doubled" in x["excerpt"].lower() and x["evidence_type"]=="direct" for x in out["evidence"])

def test_filesystem_timeout_isolated_and_index_progress_preserved(corpus, monkeypatch):
    from repo_research import extract
    original_stat = Path.stat
    broken = corpus / "cloud-timeout.pdf"
    broken.write_bytes(b"placeholder")

    def flaky_stat(self, *args, **kwargs):
        if self == broken:
            raise OSError(60, "Operation timed out", str(self))
        return original_stat(self, *args, **kwargs)

    monkeypatch.setattr(Path, "stat", flaky_stat)
    out = research(corpus, mode="inventory", depth="quick", fresh=True, backend=MockBackend())
    assert sum(out["inventory"]["file_families"].values()) >= 8
    assert len(out["inventory"]["extraction_failures"]) >= 1

def test_per_file_deadline_interrupts_stalled_io():
    from repo_research.index import _file_deadline
    with pytest.raises(TimeoutError):
        with _file_deadline(0.02):
            time.sleep(0.2)

def test_hash_subprocess_timeout_becomes_extraction_failure(corpus, monkeypatch):
    import subprocess
    from repo_research import extract
    def timed_out(*args, **kwargs):
        raise subprocess.TimeoutExpired(args[0], kwargs.get("timeout", 1))
    monkeypatch.setattr(extract.subprocess, "run", timed_out)
    chunks, failure = extract.extract_file(corpus / "primary.md", corpus,
        {"chunk_chars": 4000, "chunk_overlap_chars": 400, "file_timeout_seconds": 1})
    assert chunks == []
    assert failure and "timed out" in failure["error"].lower()

def test_stitched_quote_keeps_only_exact_contiguous_segment(corpus):
    from repo_research.index import Index
    from repo_research.research import _evidence_from_finding
    from repo_research.store import Store
    idx = Index(corpus, Store(corpus))
    idx.build({"ignore_dirs": [".codex-research"], "include": [], "exclude": [],
               "max_file_bytes": 20_000_000, "chunk_chars": 4000,
               "chunk_overlap_chars": 400, "file_timeout_seconds": 0})
    chunk = next(c for c in idx.all_chunks() if c.source_path == "primary.md")
    finding = {"excerpt": "On 15 March 2023, the Board approved establishment of the Atlas Unit... wording not present",
               "topic": "approval", "evidence_type": "direct", "stage": "approval"}
    ev = _evidence_from_finding(finding, chunk, "q", "mock", "s", chunk.text)
    assert ev and ev.excerpt.endswith("Atlas Unit")
    assert "excerpt_trimmed_to_contiguous_exact_segment" in ev.validation_flags

def test_atomic_decomposition_preserves_dates_and_stage_contrasts():
    parts=_decompose_claim("The amendment was signed on 4 May 2022, but execution began on 1 July 2022; the proposal was not implementation.")
    assert len(parts) == 3
    assert any("signed on 4 May 2022" in p for p in parts)
    assert any("execution began on 1 July 2022" in p for p in parts)

def test_duplicate_source_hashes_are_retrieved_once_and_sources_are_diverse(corpus):
    (corpus/"duplicate-primary.md").write_bytes((corpus/"primary.md").read_bytes())
    out=research(corpus,"Atlas Unit approval and operations began",mode="evidence",depth="standard",fresh=True,backend=MockBackend())
    paths_by_hash={}
    for item in out["evidence"]:
        paths_by_hash.setdefault(item["source_hash"],set()).add(item["source_path"])
    assert all(len(paths) == 1 for paths in paths_by_hash.values())
    assert out["telemetry"]["candidate_source_hashes"] == len(set(out["telemetry"]["candidate_source_hash_list"]))
    assert out["telemetry"]["candidate_files"] >= 2

def test_malformed_worker_output_is_persisted_even_after_successful_retry(corpus):
    class RetryingBackend(MockBackend):
        def __init__(self):
            super().__init__()
            self.malformed_log=[{"attempt":1,"raw":"{\"findings\":[","error":"truncated JSON","retry":True}]
    out=research(corpus,"When did Atlas become operational?",mode="evidence",depth="quick",fresh=True,backend=RetryingBackend())
    failures=[json.loads(x) for x in (corpus/".codex-research/qwen_failures.jsonl").read_text().splitlines()]
    assert failures and failures[0]["raw"] == '{"findings":['
    assert failures[0]["search_id"] == out["search_id"]
    assert out["telemetry"]["qwen_failures"] == 1
