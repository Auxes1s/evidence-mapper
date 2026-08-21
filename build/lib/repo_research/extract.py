from __future__ import annotations

import csv, hashlib, json, mimetypes, os, shutil, subprocess, tempfile
from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from .models import Chunk

TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".json", ".jsonl", ".yaml", ".yml", ".toml",
    ".xml", ".html", ".htm", ".css", ".scss", ".js", ".jsx", ".ts", ".tsx", ".py",
    ".r", ".rmd", ".qmd", ".do", ".ado", ".java", ".c", ".h", ".cpp", ".hpp", ".go",
    ".rs", ".rb", ".php", ".sh", ".zsh", ".bash", ".sql", ".tex", ".bib", ".ini",
    ".cfg", ".conf", ".env.example", ".log", ".gitignore", ".dockerfile", ".makefile",
}
SUPPORTED = TEXT_EXTENSIONS | {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".tsv"}

def file_hash(path: Path, timeout: int | float | None = None) -> str:
    # A child process gives us a hard timeout for cloud-backed reads. Python
    # signals cannot reliably interrupt a kernel filesystem call on macOS.
    shasum = shutil.which("shasum")
    if timeout and shasum:
        result = subprocess.run([shasum, "-a", "256", str(path)], capture_output=True,
                                text=True, check=True, timeout=timeout)
        return result.stdout.split()[0]
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1024 * 1024), b""): h.update(block)
    return h.hexdigest()

def stable_id(*parts: object) -> str:
    return hashlib.sha256("\x1f".join(str(x) for x in parts).encode()).hexdigest()[:24]

def _windows(text: str, size: int, overlap: int) -> Iterable[tuple[int, str]]:
    if len(text) <= size:
        yield 0, text
        return
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            split = max(text.rfind("\n\n", start, end), text.rfind("\n", start, end))
            if split > start + size // 2: end = split
        yield start, text[start:end]
        if end == len(text): break
        start = max(start + 1, end - overlap)

def _chunk_text(path: Path, rel: str, source_hash: str, text: str, cfg: dict,
                **location) -> list[Chunk]:
    out = []
    base_line = location.pop("line_start", None)
    for ordinal, (offset, part) in enumerate(_windows(text, cfg["chunk_chars"], cfg["chunk_overlap_chars"])):
        ls = base_line + text[:offset].count("\n") if base_line is not None else None
        le = ls + part.count("\n") if ls is not None else None
        loc = dict(location, line_start=ls, line_end=le)
        out.append(Chunk(stable_id(rel, source_hash, location, ordinal, offset), rel,
                         path.suffix.lower().lstrip(".") or "text", source_hash,
                         part.strip(), ordinal, **loc))
    return [x for x in out if x.text]

def _ocr_pdf(path: Path, rel: str, digest: str, page_count: int, cfg: dict) -> list[Chunk]:
    if not cfg.get("ocr_image_only_pdfs"):
        return []
    pdftoppm, tesseract = shutil.which("pdftoppm"), shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        raise RuntimeError("image-only PDF requires pdftoppm and tesseract for OCR")
    limit = min(page_count, int(cfg.get("ocr_max_pages", 120)))
    dpi = str(int(cfg.get("ocr_dpi", 200)))
    timeout = int(cfg.get("ocr_timeout_per_page_seconds", 90))
    out: list[Chunk] = []
    with tempfile.TemporaryDirectory(prefix="repo-research-ocr-") as temp:
        for page_no in range(1, limit + 1):
            stem = str(Path(temp) / f"page-{page_no}")
            subprocess.run([pdftoppm, "-f", str(page_no), "-l", str(page_no), "-r", dpi,
                            "-png", "-singlefile", str(path), stem], check=True,
                           capture_output=True, timeout=timeout)
            image = stem + ".png"
            result = subprocess.run([tesseract, image, "stdout", "-l", "eng"], check=True,
                                    capture_output=True, text=True, timeout=timeout)
            text = result.stdout.strip()
            if text:
                out += _chunk_text(path, rel, digest, text, cfg, page=page_no, section="OCR")
    return out

def extract_file(path: Path, root: Path, cfg: dict) -> tuple[list[Chunk], dict | None]:
    rel, ext, digest = str(path.relative_to(root)), path.suffix.lower(), ""
    try:
        digest = file_hash(path, cfg.get("file_timeout_seconds"))
        if ext == ".pdf":
            reader, out = PdfReader(str(path)), []
            for page_no, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                out += _chunk_text(path, rel, digest, text, cfg, page=page_no)
            if not out:
                out = _ocr_pdf(path, rel, digest, len(reader.pages), cfg)
            if not out: raise ValueError("PDF has no usable extractable text (OCR unavailable or produced no text)")
            return out, None
        if ext == ".docx":
            doc, out, heading, para = Document(str(path)), [], None, 0
            buffer, start = [], 1
            def flush():
                nonlocal buffer, start, out
                if buffer:
                    out += _chunk_text(path, rel, digest, "\n".join(buffer), cfg,
                                       section=heading, paragraph_start=start,
                                       paragraph_end=para)
                    buffer = []
            for p in doc.paragraphs:
                para += 1
                if p.style and p.style.name.startswith("Heading"):
                    flush(); heading = p.text.strip(); start = para
                elif p.text.strip():
                    if not buffer: start = para
                    buffer.append(p.text)
            flush()
            for ti, table in enumerate(doc.tables, 1):
                rows = [" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
                out += _chunk_text(path, rel, digest, "\n".join(rows), cfg, section=f"Table {ti}")
            return out, None
        if ext in {".xlsx", ".xlsm"}:
            wb, out = load_workbook(path, read_only=True, data_only=True), []
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                if not rows: continue
                headers = [str(x) if x is not None else "" for x in rows[0]]
                for start in range(1, len(rows), 40):
                    subset = rows[start:start + 40]
                    rendered = ["HEADERS: " + " | ".join(headers)]
                    for rn, row in enumerate(subset, start + 1):
                        rendered.append(f"ROW {rn}: " + " | ".join("" if x is None else str(x) for x in row))
                    out += _chunk_text(path, rel, digest, "\n".join(rendered), cfg,
                                       sheet=ws.title, row_start=start + 1,
                                       row_end=start + len(subset))
            return out, None
        if ext in {".csv", ".tsv"}:
            delimiter = "\t" if ext == ".tsv" else ","
            with path.open(encoding="utf-8-sig", errors="replace", newline="") as fh:
                rows = list(csv.reader(fh, delimiter=delimiter))
            if not rows: return [], None
            out, headers = [], rows[0]
            for start in range(1, len(rows), 60):
                rendered = ["HEADERS: " + " | ".join(headers)] + [
                    f"ROW {rn}: " + " | ".join(row) for rn, row in enumerate(rows[start:start+60], start+1)]
                out += _chunk_text(path, rel, digest, "\n".join(rendered), cfg,
                                   row_start=start + 1, row_end=min(len(rows), start + 60))
            return out, None
        text = path.read_text(encoding="utf-8", errors="replace")
        return _chunk_text(path, rel, digest, text, cfg, line_start=1), None
    except Exception as exc:
        return [], {"source_path": rel, "source_hash": digest, "error": str(exc), "source_type": ext.lstrip(".")}

def eligible_files(root: Path, cfg: dict, failures: list[dict] | None = None) -> list[Path]:
    from fnmatch import fnmatch
    out = []
    failures = failures if failures is not None else []
    ignores, includes, excludes = set(cfg["ignore_dirs"]), cfg.get("include", []), cfg.get("exclude", [])
    def walk_error(exc: OSError) -> None:
        failures.append({"source_path": getattr(exc, "filename", "") or "", "source_hash": "",
                         "error": str(exc), "source_type": "filesystem"})
    for current, dirnames, filenames in os.walk(root, onerror=walk_error):
        current_path = Path(current)
        dirnames[:] = [d for d in dirnames if d not in ignores]
        for name in filenames:
            p = current_path / name
            try:
                rel = str(p.relative_to(root))
                if p.stat().st_size > cfg["max_file_bytes"]: continue
                if includes and not any(fnmatch(rel, x) for x in includes): continue
                if any(fnmatch(rel, x) for x in excludes): continue
                if p.suffix.lower() in SUPPORTED or p.name.lower() in {"dockerfile", "makefile", "license"}:
                    out.append(p)
            except OSError as exc:
                failures.append({"source_path": str(p), "source_hash": "", "error": str(exc),
                                 "source_type": p.suffix.lower().lstrip(".") or "filesystem"})
    return sorted(out)
